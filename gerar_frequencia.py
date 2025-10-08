import calendar
import holidays
from datetime import date
from docx import Document
import os

# --- CONFIGURAÇÕES ---
# Defina o ano e o mês para o qual você quer gerar as folhas de frequência
ANO = 2026
MES = 1  # 5 para Maio, 6 para Junho, etc.

# nome dos arquivos de entrada
ARQUIVO_TEMPLATE = 'template_frequencia.docx'
ARQUIVO_FUNCIONARIO = 'funcionarios.txt'

# --- FIM DAS CONFIGURAÇÕES ---

def obter_nome_mes(mes):
    """Retorna o nome do mês em português"""
    nomes_meses = {
        1: "JANEIRO", 2: "FEVEREIRO", 3: "MARÇO", 4: "ABRIL",
        5: "MAIO", 6: "JUNHO", 7: "JULHO", 8: "AGOSTO",
        9: "SETEMBRO", 10: "OUTUBRO", 11: "NOVEMBRO", 12: "DEZEMBRO"
    }
    return nomes_meses.get(mes, "MÊS INVÁLIDO")

def substituir_placeholders(doc, placeholders):
    """Substitui os placeholders no documento (paragrafos e cabeçalhos)"""
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            if key in p.text:
                # Substitui mantendo a formatação original
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text
                        inline[i].font.bold = True # Deixa em negrito

def gerar_frequencias():
    """Função principal para gerar as folhas de frequencia."""

    # Carregar feriados nacionais e do Ceará
    feriados_br = holidays.Brazil()
    feriados_ce = holidays.Brazil(state='CE')
    feriados_ano = {**feriados_br, **feriados_ce}

    # Tenta ler os funcionários
    try:
        with open(ARQUIVO_FUNCIONARIO, 'r', encoding='utf-8') as f:
            funcionarios = [line.strip().split(';') for line in f if line.strip()]
    except FileNotFoundError:
        print(f'ERRO: Arquivo "{ARQUIVO_FUNCIONARIO}" não encontrado!')
        return

    # Criar pasta de saída se não existir
    pasta_saida = f"Frequencias_{obter_nome_mes(MES)}_{ANO}"
    if not os.path.exists(pasta_saida):
        os.makedirs(pasta_saida)

    print(f'Iniciando a geração de folhas de frequência para {obter_nome_mes(MES)} de {ANO}')

    # lOOP PARA CADA FUNCIONÁRIO
    for nome, cargo in funcionarios:
        print(f'gerando para: "{nome}"')

        doc = Document(ARQUIVO_TEMPLATE)

        # Definir os placeholders
        placeholders = {
            '{{NOME_FUNCIONARIO}}': nome,
            '{{CARGO_FUNCIONARIO}}': cargo,
            '{{MES_ANO}}': f"{obter_nome_mes(MES)}/{ANO}"
        }

        # Substituir placeholders de texto
        substituir_placeholders(doc, placeholders)

        # Acessar a primeira tabela do documento
        tabela = doc.tables[0]

        # Preencher os dias do mês
        num_dias_mes = calendar.monthrange(ANO, MES)[1]
        dias_de_semana_pt = ['SEGUNDA', 'TERÇA', 'QUARTA','QUINTA', 'SEXTA', 'SÁBADO', 'DOMINGO']

        for dia in range(1, 32):
            # A linha 0 é o cabeçalho, então os dias começam na linha 1
            linha = tabela.rows[dia]

            if dia <= num_dias_mes:
                data_atual = date(ANO, MES, dia)

                # Coluna 0: Numero do dia
                linha.cells[0].text = f"{dia:02d}"

                # Coluna 2: Decrição (Feriado, Sábado, Domingo)
                # O nome da coluna no seu exemplo é 'ENTRADA ASSINATURA'
                celula_descricao = linha.cells[2]

                if data_atual in feriados_ano:
                    # Encontou um feriado
                    celula_descricao.text = feriados_ano[data_atual].upper()
                elif data_atual.weekday() == 5: # Sábado
                    celula_descricao.text = "SÁBADO"
                elif data_atual.weekday() == 6: # domingo
                    celula_descricao.text = "DOMINGO"
                else:
                    celula_descricao.text = ""
            else:
                # Dia não existe no mês (ex: 31 de Abril), apaga a linha visualmente
                for cell in linha.cells:
                    cell.text = ""
        # 5. Salvar o novo arquivo
        nome_arquivo_saida = f"Frequencia_{obter_nome_mes(MES)}_{ANO}_{nome.split()[0]}.docx"
        caminho_saida = os.path.join(pasta_saida, nome_arquivo_saida)
        doc.save(caminho_saida)

    print("\nProcesso concluido")
    print(f"Os arquivos foram salvos na pasta: '{pasta_saida}'")

if __name__ == "__main__":
    gerar_frequencias()







